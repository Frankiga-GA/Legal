from __future__ import annotations

import io
import os
import re
import smtplib
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any
from collections import defaultdict
import time

from fastapi import Depends, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

import tempfile
try:
    import google.generativeai as genai
except ImportError:
    genai = None

import httpx
import pdfplumber
from docx import Document
from dotenv import load_dotenv



from auth import (  # noqa: E402
    CurrentUser,
    current_user,
)

BACKEND_DIR = Path(__file__).resolve().parent

# El backend SOLO lee secretos desde `backend/.env`. No debe leer el `.env`
# de la raiz porque ese archivo contiene claves publicas del frontend.
# Si necesitamos una variable en ambos lados, duplicar la entrada en los
# dos archivos .env correspondientes. Ver SECURITY.md seccion 2.
load_dotenv(BACKEND_DIR / ".env")


app = FastAPI(title="LUSTI Document Backend", version="0.2.0")


# =============================================================================
# Strip del prefijo `/api`
# =============================================================================
# En produccion (Vercel) Vercel invoca esta serverless function con la ruta
# ORIGINAL del request (ej. `/api/upload`). En local, el proxy de Vite ya
# recorta el prefijo antes de llegar aca. Para que el mismo backend funcione
# en ambos entornos sin duplicar rutas, este middleware recorta `/api/...`
# -> `/...` en el scope ASGI antes de que FastAPI haga el matching.
# =============================================================================
from starlette.middleware.base import BaseHTTPMiddleware  # noqa: E402
from starlette.requests import Request  # noqa: E402


class StripApiPrefixMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        path = request.scope.get("path", "")
        if path.startswith("/api/"):
            new_path = "/" + path[len("/api/"):]
            request.scope["path"] = new_path
            request.scope["raw_path"] = new_path.encode("latin-1")
        elif path == "/api":
            request.scope["path"] = "/"
            request.scope["raw_path"] = b"/"
        return await call_next(request)


app.add_middleware(StripApiPrefixMiddleware)


# =============================================================================
# CORS estricto
# =============================================================================
# Origenes leidos desde `ALLOWED_ORIGINS` (CSV) en backend/.env.
# En produccion, el comodin "*" NO esta permitido (incompatible con
# allow_credentials=True, ademas de ser un riesgo de seguridad).
# =============================================================================

def _parse_allowed_origins() -> list[str]:
    raw = os.getenv("ALLOWED_ORIGINS", "").strip()
    if not raw:
        # En desarrollo: solo localhost. En produccion: configurar la env.
        return ["http://localhost:5173", "http://127.0.0.1:5173"]
    origins = [item.strip() for item in raw.split(",") if item.strip()]
    if "*" in origins:
        raise RuntimeError(
            "ALLOWED_ORIGINS no puede contener '*'. Configura los origenes explicitos."
        )
    return origins


app.add_middleware(
    CORSMiddleware,
    allow_origins=_parse_allowed_origins(),
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=["Authorization", "Content-Type", "X-Request-Id"],
    expose_headers=["X-Lusti-Filename"],
    max_age=600,
)


@app.middleware("http")
async def add_security_headers(request: Request, call_next):
    response = await call_next(request)
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-XSS-Protection"] = "1; mode=block"
    response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    return response


# =============================================================================
# Utilidades de Seguridad (Validación y Rate Limiting)
# =============================================================================

_rate_limits = defaultdict(list)


def check_rate_limit(user_id: str, limit: int = 15, period: int = 60) -> None:
    """Valida si un usuario excedió el límite de peticiones (Rate Limit)."""
    now = time.time()
    user_requests = _rate_limits[user_id]
    
    # Filtrar marcas de tiempo fuera del periodo de gracia
    active_requests = [t for t in user_requests if now - t < period]
    _rate_limits[user_id] = active_requests
    
    if len(active_requests) >= limit:
        raise HTTPException(
            status_code=429,
            detail="Demasiadas peticiones. Por favor, intenta de nuevo más tarde."
        )
    _rate_limits[user_id].append(now)


def _validate_file(file: UploadFile, max_size_mb: int = 300) -> None:
    """Valida la extensión del archivo (whitelist) y su tamaño (máx 300MB)."""
    # 1. Validar extensión
    allowed_extensions = {".pdf", ".docx", ".txt", ".md", ".json", ".csv", ".jpg", ".jpeg", ".png", ".webp"}
    filename = (file.filename or "").lower()
    suffix = Path(filename).suffix
    if suffix not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Extensión de archivo '{suffix}' no permitida. Permitidos: {', '.join(allowed_extensions)}"
        )
    
    # 2. Validar tamaño por header
    max_size_bytes = max_size_mb * 1024 * 1024
    content_length = file.headers.get("content-length")
    if content_length:
        try:
            if int(content_length) > max_size_bytes:
                raise HTTPException(
                    status_code=413,
                    detail=f"El archivo excede el tamaño máximo permitido de {max_size_mb}MB."
                )
        except ValueError:
            pass
            
    # 3. Validar tamaño por descriptor en disco/memoria (en caso de que falte el header)
    try:
        file.file.seek(0, 2)
        size = file.file.tell()
        file.file.seek(0)
    except Exception:
        size = 0

    if size > max_size_bytes:
        raise HTTPException(
            status_code=413,
            detail=f"El archivo excede el tamaño máximo permitido de {max_size_mb}MB."
        )


class ChatRequest(BaseModel):
    message: str = Field(..., min_length=1)
    prompt: str | None = None
    file_name: str | None = None
    file_type: str | None = None
    file_text: str | None = None


class ChatResponse(BaseModel):
    message: str
    extracted_text: str | None = None
    file_name: str | None = None
    file_type: str | None = None


class GenerateFileRequest(BaseModel):
    message: str = Field(..., min_length=1)
    prompt: str | None = None
    file_name: str | None = None
    file_type: str | None = None
    file_text: str | None = None
    content: str | None = None
    document_type: str = "informe"
    output_format: str = Field("docx", pattern="^(docx|pdf|txt)$")


class RawAskRequest(BaseModel):
    prompt: str = Field(..., min_length=1)
    temperature: float = Field(0.25, ge=0.0, le=2.0)
    max_output_tokens: int = Field(2048, ge=64, le=8192)
    response_json: bool = False
    system_prompt: str | None = None
    history: list[dict[str, str]] | None = None
    file_name: str | None = None
    file_text: str | None = None


class RawAskResponse(BaseModel):
    text: str


class ProcessUrlRequest(BaseModel):
    url: str
    file_name: str
    file_type: str | None = None


def _extract_text_from_image(content_bytes: bytes, mime_type: str) -> str:
    import base64
    import base64
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY no esta configurada en el backend.")
    
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-3.5-flash:generateContent?key={api_key}"
    base64_image = base64.b64encode(content_bytes).decode("utf-8")
    
    prompt = (
        "Por favor, transcribe de forma precisa todo el texto visible en esta imagen. "
        "Si es una conversación, mantén el formato de diálogo. Si es un documento formal, "
        "mantén la estructura de párrafos. Responde SOLAMENTE con el texto extraído, sin comentarios adicionales."
    )
    
    payload = {
        "contents": [{
            "parts": [
                {"text": prompt},
                {"inline_data": {"mime_type": mime_type, "data": base64_image}}
            ]
        }],
        "generationConfig": {
            "temperature": 0.1,
            "maxOutputTokens": 4096,
        }
    }
    
    headers = {
        "Content-Type": "application/json",
    }
    
    with httpx.Client(timeout=90) as client:
        response = client.post(url, json=payload, headers=headers)
        
    if response.status_code >= 400:
        raise RuntimeError(f"Gemini Vision {response.status_code}: {response.text}")
        
    data = response.json()
    candidates = data.get("candidates") or []
    if not candidates:
        raise RuntimeError("Gemini Vision devolvio respuesta vacia")
        
    extracted = candidates[0].get("content", {}).get("parts", [{}])[0].get("text", "")
    return extracted.strip()


def _read_text_from_upload(upload: UploadFile) -> str:
    content = upload.file.read()
    if not content:
        return ""

    filename = (upload.filename or "").lower()
    content_type = (upload.content_type or "").lower()

    if filename.endswith((".txt", ".md", ".csv", ".json")) or content_type.startswith("text/"):
        return content.decode("utf-8", errors="ignore").strip()

    if filename.endswith((".jpg", ".jpeg", ".png", ".webp")) or content_type.startswith("image/"):
        try:
            mime = content_type if content_type.startswith("image/") else "image/jpeg"
            return _extract_text_from_image(content, mime)
        except Exception as e:
            print(f"Error OCR Vision: {e}")
            raise RuntimeError(f"No se pudo extraer texto de la imagen: {e}")

    if filename.endswith(".pdf") or content_type == "application/pdf":
        # 1. Intentar pdfplumber (limitado a las primeras 25 páginas para evitar OOM/Timeout en Vercel Serverless)
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages[:25]]
            text = "\n\n".join(part.strip() for part in pages if part.strip()).strip()
            if text:
                return text
        except Exception as e:
            print(f"pdfplumber exception: {e}")



        return ""

    if filename.endswith(".docx") or content_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }:
        try:
            doc = Document(io.BytesIO(content))
            paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
            return "\n".join(paragraphs).strip()
        except Exception:
            return ""

    return content.decode("utf-8", errors="ignore").strip()


def _build_document_prompt(message: str, prompt: str, file_name: str | None, file_type: str | None, file_text: str | None) -> str:
    clean_file_text = (file_text or "").strip()
    return f"""
Eres LUSTI, un asistente documental legal.
Responde en espanol claro, profesional y util.
No uses Markdown.
No inventes informacion.
Si el texto adjunto contiene campos, variables o placeholders como {{nombre}}, {{dni}} o similares, tratalo como una plantilla valida.
No digas que el documento esta vacio si hay texto real, aunque sea una plantilla corta o tenga campos por completar.
Si el documento es una plantilla, primero identifica su estructura y luego explica que datos faltan para completarla.
Si el texto adjunto realmente esta vacio o no se pudo leer, dilo claramente y pide los datos faltantes.

Prompt del asistente:
{prompt or 'Sin prompt adicional.'}

Archivo adjunto:
Nombre: {file_name or 'Sin archivo'}
Tipo: {file_type or 'Desconocido'}
Texto extraido:
{clean_file_text or 'No se pudo extraer texto util del archivo.'}

Instruccion del usuario:
{message}

Respuesta esperada:
Organiza la respuesta segun lo que pida el usuario.
Si pide resumen o analisis de documento, usa este orden:
1. Resumen breve
2. Datos clave encontrados
3. Riesgos o puntos de atencion
4. Informacion faltante
5. Siguientes pasos recomendados

Si pide extraer datos, usa este orden:
1. Partes
2. Fechas
3. Montos o remuneracion
4. Obligaciones principales
5. Observaciones

Si pide completar una plantilla:
1. Identifica la plantilla
2. Lista los campos detectados
3. Indica que campos faltan
4. Entrega el borrador completado cuando haya datos suficientes

Reglas finales:
- Lee el texto adjunto.
- No cortes frases ni dejes campos a medias.
- Si el documento es largo, prioriza una respuesta resumida pero cerrada.
- Si falta informacion, indica exactamente que falta.
""".strip()


async def _ask_gemini(
    prompt_text: str,
    *,
    temperature: float = 0.25,
    max_output_tokens: int = 4096,
    response_json: bool = False,
    system_prompt: str | None = None,
    history: list[dict[str, str]] | None = None,
    file_name: str | None = None,
    file_text: str | None = None,
) -> str:
    """Llama a Groq (OpenAI-compatible) usando GROQ_API_KEY.

    Si se pasa `history`, se construye la conversacion como
    [system, ...history..., user(prompt_text)]. Esto le da a la IA
    memoria de los mensajes previos del chat.

    Si se pasa `file_name` y `file_text`, el contenido del archivo se
    adjunta al prompt del usuario como contexto de solo lectura.
    """
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        raise RuntimeError("GROQ_API_KEY no esta configurada en el backend.")

    model = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
    url = "https://api.groq.com/openai/v1/chat/completions"

    # Si hay archivo adjunto, lo agregamos como contexto antes de la pregunta.
    # Truncamos para no pasar el limite de tokens del modelo.
    if file_text and file_text.strip():
        max_file_chars = 18000
        truncated = file_text.strip()
        if len(truncated) > max_file_chars:
            truncated = truncated[:max_file_chars] + "\n\n[...texto truncado por limite de contexto...]"
        file_block = (
            f"Archivo adjunto: {file_name or 'documento'}\n"
            f"--- INICIO DEL ARCHIVO ---\n"
            f"{truncated}\n"
            f"--- FIN DEL ARCHIVO ---\n\n"
        )
        prompt_text = file_block + prompt_text

    messages: list[dict[str, str]] = []
    if system_prompt and system_prompt.strip():
        messages.append({"role": "system", "content": system_prompt.strip()})

    if history:
        # Solo aceptamos roles validos; descartamos entradas mal formadas
        for entry in history:
            role = entry.get("role") if isinstance(entry, dict) else None
            content = entry.get("content") if isinstance(entry, dict) else None
            if role in ("user", "ai", "assistant", "system") and content:
                # Groq espera "assistant", no "ai"
                normalized = "assistant" if role == "ai" else role
                messages.append({"role": normalized, "content": str(content)})

    messages.append({"role": "user", "content": prompt_text})

    payload: dict[str, Any] = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_output_tokens,
        "top_p": 0.9,
    }
    if response_json:
        payload["response_format"] = {"type": "json_object"}

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    async with httpx.AsyncClient(timeout=90) as client:
        response = await client.post(url, json=payload, headers=headers)

    if response.status_code >= 400:
        raise RuntimeError(f"Groq {response.status_code}: {response.text}")

    data = response.json()
    choices = data.get("choices") or []
    if not choices:
        raise RuntimeError(f"Groq devolvio respuesta vacia: {data}")
    extracted = (choices[0].get("message") or {}).get("content") or ""
    extracted = extracted.strip()
    if not extracted:
        raise RuntimeError("Groq devolvio contenido vacio")
    return extracted


def _build_file_generation_prompt(payload: GenerateFileRequest) -> str:
    clean_file_text = (payload.file_text or "").strip()
    return f"""
Eres LUSTI, un redactor legal para estudios juridicos peruanos.
Genera el contenido final de un documento legal en espanol claro, profesional y editable.
No uses Markdown ni asteriscos.
No inventes datos. Si falta informacion, deja campos entre corchetes como [DATO PENDIENTE].

Tipo de documento solicitado:
{payload.document_type or 'informe'}

Instruccion del usuario:
{payload.message}

Prompt adicional:
{payload.prompt or 'Sin prompt adicional.'}

Archivo base:
Nombre: {payload.file_name or 'Sin archivo'}
Tipo: {payload.file_type or 'Desconocido'}
Texto extraido:
{clean_file_text or 'Sin texto extraido.'}

Estructura recomendada:
Titulo
Fecha
Antecedentes
Analisis
Riesgos o puntos de atencion
Solicitud, conclusion o recomendacion
Datos faltantes

Entrega solo el contenido del documento final.
""".strip()


def _clean_filename(value: str, extension: str) -> str:
    base = re.sub(r"[^a-zA-Z0-9_-]+", "-", value.strip().lower()).strip("-")
    if not base:
        base = "documento-lusti"
    return f"{base[:80]}.{extension}"


def _split_document_blocks(content: str) -> list[str]:
    blocks = [part.strip() for part in re.split(r"\n{2,}", content.strip()) if part.strip()]
    if blocks:
        return blocks
    return [line.strip() for line in content.splitlines() if line.strip()]


def _clean_final_document_content(content: str) -> str:
    def plain(value: str) -> str:
        normalized_value = unicodedata.normalize("NFKD", value)
        return "".join(char for char in normalized_value if not unicodedata.combining(char)).lower()

    lines = [line.strip() for line in content.replace("\r\n", "\n").replace("\r", "\n").splitlines()]
    cleaned: list[str] = []
    skip_until_blank = False
    boilerplate_patterns = [
        r"^estimad[oa]\s+(usuario|cliente|senor|senora)",
        r"^he procesado\b",
        r"^a continuacion\b.*\b(documento|pdf|docx|contenido)\b",
        r"^a continuaci.n\b.*\b(documento|pdf|docx|contenido)\b",
        r"^este (documento|analisis|an.lisis|informe) (es|tiene) .*preliminar",
        r"^no constituye asesoria legal",
        r"^no constituye asesor.a legal",
        r"^se recomienda la revision",
        r"^se recomienda la revisi.n",
        r"^por favor,?\s+proporcione\b",
        r"^¿?le gustaria\b",
        r"^\??le gustar.a\b",
    ]

    for line in lines:
        normalized = plain(line).strip(" .:")
        if skip_until_blank:
            if not line:
                skip_until_blank = False
            continue
        if not line:
            if cleaned and cleaned[-1] != "":
                cleaned.append("")
            continue
        if line in {"---", "***"}:
            if cleaned and cleaned[-1] != "":
                cleaned.append("")
            continue
        if re.search(r"^informaci.n faltante", normalized):
            skip_until_blank = True
            continue
        if normalized.startswith("siguientes pasos recomendados") or normalized.startswith("siguientes pasos"):
            skip_until_blank = True
            continue
        if any(re.search(pattern, normalized, flags=re.IGNORECASE) for pattern in boilerplate_patterns):
            continue
        cleaned.append(line)

    result = "\n".join(cleaned).strip()
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result or content.strip()


def _derive_document_title(content: str, fallback: str) -> tuple[str, str]:
    blocks = _split_document_blocks(content)
    if not blocks:
        return fallback, content

    first = blocks[0].strip().strip("#").strip()
    is_title = (
        4 <= len(first) <= 90
        and not first.endswith(".")
        and not first.lower().startswith(("nombre:", "dni:", "fecha:", "cargo:", "para:", "de:"))
    )
    if not is_title:
        return fallback, content

    remaining = content.strip()[len(blocks[0]):].lstrip()
    return first, remaining


def _build_docx_bytes(content: str, title: str) -> bytes:
    document = Document()
    document.add_heading(title, level=1)

    for block in _split_document_blocks(content):
        clean_block = block.strip()
        if not clean_block:
            continue

        is_heading = (
            len(clean_block) <= 90
            and not clean_block.endswith(".")
            and not clean_block.startswith(("-", "1.", "2.", "3.", "4.", "5."))
        )
        if is_heading:
            document.add_heading(clean_block.rstrip(":"), level=2)
        else:
            for line in clean_block.splitlines():
                stripped = line.strip()
                if stripped:
                    document.add_paragraph(stripped)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_pdf_bytes(content: str, title: str) -> bytes:
    raise NotImplementedError("PDF generation on backend is disabled to save Vercel bundle size. Use frontend export.")


def _build_txt_bytes(content: str) -> bytes:
    return content.encode("utf-8")


@app.get("/health")
def health() -> dict[str, Any]:
    return {
        "status": "ok",
        "ts": datetime.utcnow().isoformat() + "Z",
        "groq_configured": "yes" if os.getenv("GROQ_API_KEY") else "no",
        "groq_model": os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile"),
        "auth_required": "yes" if os.getenv("SUPABASE_JWT_SECRET") else "no",
        "version": app.version,
    }


# =============================================================================
# Proxy a El Peruano (evita CORS en el navegador)
# =============================================================================
# El frontend no puede hacer fetch directo a elperuano.pe porque el sitio no
# envia Access-Control-Allow-Origin. Pasamos por aca, el backend no tiene
# esa restriccion.
EL_PERUANO_SEARCH_URL = "https://www.elperuano.pe/portal/buscador"
EL_PERUANO_API_URL = "https://www.elperuano.pe/portal/_SearchNews"
EL_PERUANO_USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/124.0 Safari/537.36"
)


def _normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").strip())


def _infer_category(text: str) -> str:
    n = text.lower()
    if "trabajo" in n or "laboral" in n or "sunafil" in n:
        return "Laboral"
    if "sunat" in n or "tribut" in n:
        return "Tributario"
    if "indecopi" in n or "consumidor" in n:
        return "Consumidor"
    if "contratacion" in n or "osce" in n:
        return "Contrataciones"
    if "empresa" in n or "societ" in n:
        return "Corporativo"
    if "penal" in n or "delito" in n:
        return "Penal"
    if "civil" in n or "propiedad" in n:
        return "Civil"
    return "General"


def _infer_type(title: str) -> str:
    n = title.lower()
    if "decreto supremo" in n:
        return "Decreto Supremo"
    if "decreto legislativo" in n:
        return "Decreto Legislativo"
    if "resolucion" in n or "resolución" in n:
        return "Resolucion"
    if "ley " in n:
        return "Ley"
    if "sentencia" in n or "casacion" in n or "casación" in n:
        return "Jurisprudencia"
    return "Norma Legal"


def _build_impact(category: str) -> str:
    label = "el estudio" if category == "General" else f"la practica {category.lower()}"
    return f"Revisar posible impacto en {label} y validar si corresponde vincularla a expedientes activos."


def _parse_elperuano_dotnet_date(value: str | None) -> str:
    """Convierte /Date(1773982800000)/ a YYYY-MM-DD."""
    if not value:
        return datetime.utcnow().date().isoformat()
    match = re.search(r"/Date\((-?\d+)(?:[+-]\d+)?\)/", value)
    if not match:
        return datetime.utcnow().date().isoformat()
    try:
        ts_ms = int(match.group(1))
        return datetime.utcfromtimestamp(ts_ms / 1000).date().isoformat()
    except Exception:
        return datetime.utcnow().date().isoformat()


def _map_seccion_to_category(seccion: str, title: str) -> str:
    text = f"{seccion or ''} {title or ''}".lower()
    if "derecho" in text or "legal" in text:
        # Si no hay pista adicional, devolvemos General para no inventar materia
        return _infer_category(title)
    return _infer_category(text)


@app.get("/elperuano/search")
async def elperuano_search(
    q: str = "",
    pageSize: int = 10,
    _: CurrentUser = Depends(current_user),
) -> dict[str, Any]:
    """Proxy que consulta el buscador publico de El Peruano desde el servidor.

    Acepta `?q=termino` para hacer una busqueda real (el endpoint oficial es
    /portal/_SearchNews?claves=...). Si `q` viene vacio, devuelve los registros
    mas recientes de la portada.
    """
    query = (q or "").strip()
    page_size = max(1, min(int(pageSize or 10), 25))

    headers = {
        "User-Agent": EL_PERUANO_USER_AGENT,
        "Accept": "application/json, text/plain, */*",
        "Referer": EL_PERUANO_SEARCH_URL,
    }
    params = {"pageIndex": "1", "pageSize": str(page_size), "claves": query}

    try:
        async with httpx.AsyncClient(timeout=20.0, follow_redirects=True, headers=headers) as client:
            response = await client.get(EL_PERUANO_API_URL, params=params)
            response.raise_for_status()
            payload = response.json()
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"No se pudo consultar El Peruano: {exc}") from exc

    if not isinstance(payload, list):
        return {
            "items": [],
            "source": "empty",
            "checkedAt": datetime.utcnow().isoformat() + "Z",
            "query": query,
            "error": None,
        }

    items: list[dict[str, Any]] = []
    for index, raw in enumerate(payload):
        title = _normalize_text(str(raw.get("vchTitulo") or ""))
        if not title or len(title) < 8:
            continue
        url_path = str(raw.get("URLFriendLy") or "").lstrip("./").lstrip("/")
        full_url = f"https://www.elperuano.pe/{url_path}" if url_path else "#"
        seccion = _normalize_text(str(raw.get("Seccion") or ""))
        date = _parse_elperuano_dotnet_date(raw.get("dtmFecha"))
        description = _normalize_text(str(raw.get("vchDescripcion") or ""))
        category = _map_seccion_to_category(seccion, title)
        items.append({
            "id": f"elperuano-{query or 'recientes'}-{index}-{date}",
            "title": title,
            "date": date,
            "type": _infer_type(title),
            "source": "El Peruano",
            "entity": seccion or "Diario Oficial El Peruano",
            "summary": description or "Resultado obtenido desde el buscador publico de El Peruano. Abre la fuente para revisar el texto completo.",
            "impact": _build_impact(category),
            "url": full_url,
            "category": category,
            "urgency": "Alta" if index < 2 else "Media",
            "scrapedAt": f"Busqueda: {query}" if query else "Consulta en vivo",
            "official": True,
        })

    return {
        "items": items,
        "source": "live" if items else "empty",
        "checkedAt": datetime.utcnow().isoformat() + "Z",
        "query": query,
        "error": None,
    }


@app.get("/me")
def me(user: CurrentUser = Depends(current_user)) -> dict[str, Any]:
    return {
        "user_id": user.user_id,
        "email": user.email,
        "role": user.role,
    }


# =============================================================================
# Health Checks (para monitoreo interno)
# =============================================================================




@app.get("/health/groq")
async def health_groq() -> dict[str, Any]:
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return {"ok": False, "error": "GROQ_API_KEY no configurada"}
    try:
        async with httpx.AsyncClient(timeout=5) as client:
            resp = await client.get(
                "https://api.groq.com/openai/v1/models",
                headers={"Authorization": f"Bearer {api_key}"},
            )
        return {"ok": resp.status_code == 200, "status": resp.status_code}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/health/supabase")
async def health_supabase() -> dict[str, Any]:
    url = os.getenv("SUPABASE_URL")
    service_key = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
    if not url or not service_key:
        return {"ok": False, "error": "Supabase no configurado"}
    try:
        async with httpx.AsyncClient(timeout=5) as client:
            resp = await client.get(
                f"{url.rstrip('/')}/rest/v1/cases?select=id&limit=1",
                headers={
                    "apikey": service_key,
                    "Authorization": f"Bearer {service_key}",
                },
            )
        return {"ok": resp.status_code == 200, "status": resp.status_code}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/health/email")
async def health_email() -> dict[str, Any]:
    gmail_user = os.getenv("GMAIL_USER")
    gmail_pass = os.getenv("GMAIL_APP_PASSWORD")
    if not gmail_user or not gmail_pass:
        return {"ok": False, "error": "SMTP no configurado"}
    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=5) as server:
            server.login(gmail_user, gmail_pass)
        return {"ok": True}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.post("/upload")
def upload_file(
    file: UploadFile = File(...),
    user: CurrentUser = Depends(current_user),
) -> dict[str, str]:
    try:
        _validate_file(file)
        extracted_text = _read_text_from_upload(file)
    except Exception as e:
        print(f"Error processing file in upload_file: {e}")
        extracted_text = ""

    return {
        "file_name": file.filename or "",
        "file_type": file.content_type or "application/octet-stream",
        "extracted_text": extracted_text,
    }



@app.post("/process-url")
async def process_url(
    payload: ProcessUrlRequest,
    user: CurrentUser = Depends(current_user),
) -> dict[str, str]:
    check_rate_limit(user.user_id, limit=5, period=60)
    
    if genai is None:
        raise HTTPException(status_code=501, detail="El servicio de Gemini Vision (google-generativeai) no está disponible en este entorno.")

    
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="GEMINI_API_KEY no configurada en el servidor.")
        
    genai.configure(api_key=api_key)
    
    # 1. Download file from Supabase URL to a temp file
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp_path = tmp.name
            async with httpx.AsyncClient(timeout=60) as client:
                async with client.stream('GET', payload.url) as response:
                    if response.status_code != 200:
                        raise HTTPException(status_code=400, detail="No se pudo descargar el archivo.")
                    async for chunk in response.aiter_bytes():
                        tmp.write(chunk)
    except Exception as e:
        print(f"Error downloading file: {e}")
        raise HTTPException(status_code=500, detail="Error descargando el archivo interno.")
        
    # 2. Upload to Gemini File API and process
    try:
        uploaded_file = genai.upload_file(path=tmp_path, display_name=payload.file_name)
        
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-3.5-flash')
        prompt = "Transcribe todo el texto visible de este documento de forma precisa. No añadas comentarios ni resúmenes, solo devuelve el texto extraído."
        response = model.generate_content([uploaded_file, prompt])
        
        extracted_text = response.text
        
        # Limpieza manual
        try:
            os.remove(tmp_path)
            genai.delete_file(uploaded_file.name)
        except Exception:
            pass
            
        return {
            "file_name": payload.file_name,
            "file_type": payload.file_type or "application/pdf",
            "extracted_text": extracted_text,
        }
    except Exception as e:
        print(f"Error con Gemini API: {e}")
        try:
            os.remove(tmp_path)
        except Exception:
            pass
        raise HTTPException(status_code=500, detail="Error procesando el documento con Gemini.")


@app.post("/ai/raw", response_model=RawAskResponse)
async def ai_raw(
    payload: RawAskRequest,
    user: CurrentUser = Depends(current_user),
) -> RawAskResponse:
    """Pasa un prompt completo a Gemini sin template del backend. Usado por el frontend."""
    check_rate_limit(user.user_id, limit=15, period=60)
    try:
        text = await _ask_gemini(
            payload.prompt,
            temperature=payload.temperature,
            max_output_tokens=payload.max_output_tokens,
            response_json=payload.response_json,
            system_prompt=payload.system_prompt,
            history=payload.history,
            file_name=payload.file_name,
            file_text=payload.file_text,
        )
    except Exception as error:
        print(f"Error calling Groq in ai_raw: {error}")
        raise HTTPException(status_code=502, detail="No se pudo procesar la solicitud con el servicio de IA.")
    return RawAskResponse(text=text)


@app.post("/chat", response_model=ChatResponse)
async def chat(
    payload: ChatRequest,
    user: CurrentUser = Depends(current_user),
) -> ChatResponse:
    check_rate_limit(user.user_id, limit=20, period=60)
    if not payload.message.strip():
        raise HTTPException(status_code=400, detail="message is required")

    prompt_text = _build_document_prompt(
        message=payload.message,
        prompt=payload.prompt or "",
        file_name=payload.file_name,
        file_type=payload.file_type,
        file_text=payload.file_text,
    )

    try:
        response_text = await _ask_gemini(prompt_text)
    except Exception as error:
        print(f"Error calling Groq in chat: {error}")
        response_text = "No se pudo procesar la solicitud con el servicio de IA en este momento."

    return ChatResponse(
        message=response_text,
        extracted_text=payload.file_text,
        file_name=payload.file_name,
        file_type=payload.file_type,
    )


@app.post("/generate-document")
def generate_document(
    message: str = Form(...),
    prompt: str = Form(""),
    file: UploadFile | None = File(None),
    user: CurrentUser = Depends(current_user),
):
    check_rate_limit(user.user_id, limit=15, period=60)
    extracted_text = ""
    file_name = None
    file_type = None

    if file is not None:
        _validate_file(file)
        file_name = file.filename
        file_type = file.content_type
        extracted_text = _read_text_from_upload(file)

    return {
        "message": message,
        "prompt": prompt,
        "file_name": file_name,
        "file_type": file_type,
        "extracted_text": extracted_text,
        "status": "ready",
    }


@app.post("/generate-file")
async def generate_file(
    payload: GenerateFileRequest,
    user: CurrentUser = Depends(current_user),
):
    check_rate_limit(user.user_id, limit=15, period=60)
    output_format = payload.output_format.lower()
    document_type = (payload.document_type or "documento").strip()

    content = (payload.content or "").strip()
    if not content:
        prompt_text = _build_file_generation_prompt(payload)
        try:
            content = await _ask_gemini(prompt_text)
        except Exception as error:
            print(f"Error calling Groq in generate_file: {error}")
            raise HTTPException(status_code=502, detail="No se pudo procesar la solicitud con el servicio de IA.")

    if not content:
        raise HTTPException(status_code=400, detail="No hay contenido para generar el archivo")

    content = _clean_final_document_content(content)
    title, content = _derive_document_title(content, f"{document_type.title()} LUSTI")

    if output_format == "docx":
        file_bytes = _build_docx_bytes(content, title)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif output_format == "pdf":
        try:
            file_bytes = _build_pdf_bytes(content, title)
        except Exception as error:
            raise HTTPException(status_code=500, detail=f"No se pudo generar PDF: {error}") from error
        media_type = "application/pdf"
    else:
        file_bytes = _build_txt_bytes(content)
        media_type = "text/plain; charset=utf-8"

    filename = _clean_filename(f"{document_type}-lusti", output_format)
    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"',
        "X-Lusti-Filename": filename,
    }

    return StreamingResponse(io.BytesIO(file_bytes), media_type=media_type, headers=headers)
